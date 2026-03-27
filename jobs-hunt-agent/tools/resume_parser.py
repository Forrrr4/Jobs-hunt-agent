"""
简历读取、解析与输出工具。

功能：
  - 从 Markdown 文件读取基础简历
  - 提取关键事实元素（公司名、日期、数字），用于完整性校验
  - 将定制后的 Markdown 简历导出为 PDF（fpdf2）
  - 将定制后的 Markdown 简历导出为 DOCX（python-docx，可选备用）
"""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Markdown 读取
# ---------------------------------------------------------------------------


def read_resume(path: str | Path) -> str:
    """
    读取 Markdown 格式的基础简历文件。

    Args:
        path: 简历文件路径（相对或绝对）

    Returns:
        简历全文字符串

    Raises:
        FileNotFoundError: 文件不存在
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"简历文件不存在：{p.resolve()}")
    text = p.read_text(encoding="utf-8")
    logger.info("已读取简历：%s（%d 字符）", p.resolve(), len(text))
    return text


# ---------------------------------------------------------------------------
# 关键事实提取（用于完整性校验）
# ---------------------------------------------------------------------------


def extract_facts(resume_text: str) -> dict:
    """
    从简历文本中提取关键事实元素，用于校验 LLM 未篡改内容。

    提取规则：
      - dates:    年份、时间段（如 2021.09、2024.07 — 2024.09）
      - numbers:  百分比和数字量化数据（如 30%、50%、10 万）
      - names:    加粗文本（通常是公司名、职位名、学校名）

    Returns:
        {
            "dates": [...],
            "numbers": [...],
            "bold_items": [...],
        }
    """
    # 时间段：限定年份前缀 19xx/20xx，避免误匹配电话号码等任意4位数字
    dates = re.findall(
        r"(?:19|20)\d{2}[.\-/年]\d{1,2}"           # 2024.01 / 2021年09
        r"|(?:19|20)\d{2}(?=\s*[—\-–])"            # 范围起始 2024 —
        r"|\b(?:19|20)\d{2}\b",                     # 独立年份 2024
        resume_text,
    )

    # 量化数据：百分比、倍数、数量（支持中文单位）
    numbers = re.findall(
        r"\d+(?:\.\d+)?(?:%|倍|万|千|百|个|条|次|人|项|MB|GB|QPS|TPS|ms)",
        resume_text,
    )

    # 加粗文本（Markdown **...** 或 __...__），通常是公司、职位、学校名
    bold_items = re.findall(r"\*\*(.+?)\*\*|__(.+?)__", resume_text)
    bold_items = [b[0] or b[1] for b in bold_items if b[0] or b[1]]

    return {
        "dates": sorted(set(dates)),
        "numbers": sorted(set(numbers)),
        "bold_items": sorted(set(bold_items)),
    }


def validate_resume_integrity(original: str, tailored: str) -> tuple[bool, list[str]]:
    """
    校验定制简历是否保留了原始简历的所有关键事实。

    逐一检查原始简历中的日期、数字、加粗项是否在定制版中仍然存在。

    Returns:
        (is_valid, violations)
        - is_valid:   True 表示完整性校验通过
        - violations: 被删除/修改的事实列表（空表示无违规）
    """
    original_facts = extract_facts(original)
    tailored_facts_raw = extract_facts(tailored)

    # 把定制版全文也当字符串匹配（避免格式变化影响正则）
    violations: list[str] = []

    # 检查日期
    for date in original_facts["dates"]:
        if date not in tailored:
            violations.append(f"[日期丢失] {date!r}")

    # 检查量化数据
    for num in original_facts["numbers"]:
        if num not in tailored:
            violations.append(f"[数字丢失] {num!r}")

    # 检查加粗项（公司名、学校名等）——允许格式稍有变化，用子串匹配
    for item in original_facts["bold_items"]:
        # 某些项目名可能很短且通用，跳过 3 字以下的
        if len(item) < 4:
            continue
        if item not in tailored:
            violations.append(f"[加粗项丢失] {item!r}")

    is_valid = len(violations) == 0
    if not is_valid:
        logger.warning("简历完整性校验失败（%d 项违规）：%s", len(violations), violations)
    else:
        logger.info("简历完整性校验通过")

    return is_valid, violations


# ---------------------------------------------------------------------------
# 输出：保存 Markdown
# ---------------------------------------------------------------------------


def save_resume_markdown(
    content: str,
    output_dir: str | Path,
    job_id: str,
    date_str: Optional[str] = None,
) -> Path:
    """
    将定制简历保存为 Markdown 文件。

    文件名格式：{job_id}_{YYYYMMDD}.md

    Returns:
        保存路径
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    date_str = date_str or datetime.now().strftime("%Y%m%d")
    # job_id 可能含特殊字符，做简单清理
    safe_id = re.sub(r"[^\w\-]", "_", job_id)
    filename = f"{safe_id}_{date_str}.md"
    path = output_dir / filename
    path.write_text(content, encoding="utf-8")
    logger.info("简历 Markdown 已保存：%s", path)
    return path


# ---------------------------------------------------------------------------
# 输出：Markdown → PDF（fpdf2 后端，纯 Python，无系统依赖）
# ---------------------------------------------------------------------------


def save_resume_pdf(
    markdown_text: str,
    output_dir: str | Path,
    job_id: str,
    date_str: Optional[str] = None,
) -> Optional[Path]:
    """
    将 Markdown 简历渲染为 PDF 文件（fpdf2 后端）。

    渲染策略：
      - 解析 Markdown 结构（标题/列表/加粗/正文），逐段写入 PDF
      - 支持中文（使用 fpdf2 内置 Unicode 支持）
      - 不依赖任何系统字体或 GTK 库

    Args:
        markdown_text: 完整简历 Markdown 文本
        output_dir:    输出目录
        job_id:        职位 ID（用于文件名）
        date_str:      日期字符串（YYYYMMDD），默认当天

    Returns:
        PDF 文件路径，失败时返回 None
    """
    try:
        from fpdf import FPDF
    except ImportError:
        logger.error("fpdf2 未安装，无法生成 PDF。请运行：pip install fpdf2")
        return None

    try:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        date_str = date_str or datetime.now().strftime("%Y%m%d")
        safe_id = re.sub(r"[^\w\-]", "_", job_id)
        pdf_path = output_dir / f"{safe_id}_{date_str}.pdf"

        pdf = _build_pdf(markdown_text)
        pdf.output(str(pdf_path))
        logger.info("简历 PDF 已保存：%s", pdf_path)
        return pdf_path

    except Exception as e:
        logger.error("PDF 生成失败：%s", e)
        return None


def _build_pdf(markdown_text: str):
    """内部：将 Markdown 文本转为 FPDF 对象（支持中文）。"""
    from fpdf import FPDF

    # 查找系统中文字体（Windows 优先使用 SimHei）
    _CJK_FONT_CANDIDATES = [
        Path("C:/Windows/Fonts/simhei.ttf"),   # Windows 黑体（单 TTF，最易加载）
        Path("C:/Windows/Fonts/simkai.ttf"),   # Windows 楷体
        Path("C:/Windows/Fonts/simsun.ttc"),   # Windows 宋体（TTC 格式）
        Path("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),  # Linux
        Path("/System/Library/Fonts/PingFang.ttc"),               # macOS
    ]
    cjk_font_path: Optional[Path] = None
    for candidate in _CJK_FONT_CANDIDATES:
        if candidate.exists():
            cjk_font_path = candidate
            break

    class ResumePDF(FPDF):
        def header(self):
            pass

        def footer(self):
            self.set_y(-12)
            if cjk_font_path:
                self.set_font("CJK", size=8)
            else:
                self.set_font("Helvetica", "I", 8)
            self.set_text_color(150, 150, 150)
            self.cell(0, 8, f"Page {self.page_no()}", align="C")

    pdf = ResumePDF(orientation="P", unit="mm", format="A4")
    pdf.set_margins(left=20, top=20, right=20)
    pdf.set_auto_page_break(auto=True, margin=15)

    # 注册中文字体
    if cjk_font_path:
        pdf.add_font("CJK", fname=str(cjk_font_path))
        pdf.add_font("CJK", style="B", fname=str(cjk_font_path))  # 粗体用同一字体模拟
        logger.debug("PDF 使用中文字体：%s", cjk_font_path.name)
    else:
        logger.warning("未找到中文字体，PDF 中文字符将显示为问号。建议安装 SimHei/WQY 字体")

    pdf._cjk_available = cjk_font_path is not None
    pdf.add_page()
    _render_markdown(pdf, markdown_text)
    return pdf


def _render_markdown(pdf, text: str) -> None:
    """
    按行解析 Markdown，逐段写入 FPDF。

    支持：
      # H1（姓名/大标题）
      ## H2（章节标题）
      ### H3（项目/经历小标题）
      - 列表项
      **加粗** 文本
      --- 分隔线
      普通段落
    """
    # 颜色常量
    COLOR_H1 = (20, 80, 160)
    COLOR_H2 = (50, 50, 50)
    COLOR_H3 = (70, 70, 70)
    COLOR_BODY = (60, 60, 60)
    COLOR_LINE = (180, 180, 180)

    cjk = getattr(pdf, "_cjk_available", False)

    def fnt(bold: bool = False, italic: bool = False, size: int = 10):
        """统一字体选择：有中文字体用 CJK，否则退回 Helvetica。"""
        if cjk:
            style = "B" if bold else ""
            pdf.set_font("CJK", style=style, size=size)
        else:
            style = ("B" if bold else "") + ("I" if italic else "")
            pdf.set_font("Helvetica", style=style, size=size)

    def set_color(rgb):
        pdf.set_text_color(*rgb)

    def draw_line():
        pdf.set_draw_color(*COLOR_LINE)
        pdf.line(20, pdf.get_y(), 190, pdf.get_y())
        pdf.ln(3)

    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        i += 1

        # 空行
        if not line.strip():
            pdf.ln(2)
            continue

        # 分隔线
        if re.match(r"^-{3,}$", line.strip()) or re.match(r"^={3,}$", line.strip()):
            draw_line()
            continue

        # H1：# 标题
        if line.startswith("# "):
            content = line[2:].strip()
            fnt(bold=True, size=20)
            set_color(COLOR_H1)
            pdf.multi_cell(0, 10, _strip_md(content), align="C")
            pdf.ln(1)
            draw_line()
            continue

        # H2：## 标题
        if line.startswith("## "):
            content = line[3:].strip()
            pdf.ln(3)
            fnt(bold=True, size=13)
            set_color(COLOR_H2)
            pdf.multi_cell(0, 7, _strip_md(content))
            # 下划线
            y = pdf.get_y()
            pdf.set_draw_color(*COLOR_H2)
            pdf.line(20, y, 190, y)
            pdf.ln(3)
            continue

        # H3：### 标题
        if line.startswith("### "):
            content = line[4:].strip()
            pdf.ln(2)
            fnt(bold=True, size=11)
            set_color(COLOR_H3)
            pdf.multi_cell(0, 6, _strip_md(content))
            pdf.ln(1)
            continue

        # 列表项：- 或 * 开头
        if re.match(r"^[\-\*]\s+", line):
            content = re.sub(r"^[\-\*]\s+", "", line)
            fnt(size=10)
            set_color(COLOR_BODY)
            # 缩进 + 项目符号
            pdf.set_x(25)
            pdf.cell(5, 6, "-")  # bullet（CJK 字体不含 U+2022，用 ASCII 横杠替代）
            pdf.set_x(30)
            pdf.multi_cell(160, 6, _strip_md(content))
            continue

        # 引用块：> 开头（如简历顶部的说明）
        if line.startswith("> "):
            content = line[2:].strip()
            fnt(italic=True, size=9)
            set_color((120, 120, 120))
            pdf.set_x(25)
            pdf.multi_cell(165, 5, _strip_md(content))
            continue

        # 普通段落
        fnt(size=10)
        set_color(COLOR_BODY)
        pdf.multi_cell(0, 6, _strip_md(line))


def _strip_md(text: str) -> str:
    """
    去除内联 Markdown 标记（加粗、斜体、链接、代码），保留纯文本。
    fpdf2 不支持富文本，统一转为普通文本输出。
    """
    # **bold** or __bold__
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    # *italic* or _italic_
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    # `code`
    text = re.sub(r"`(.+?)`", r"\1", text)
    # [link](url)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    # 去除多余空白
    return text.strip()


# ---------------------------------------------------------------------------
# 输出：Markdown → DOCX（python-docx，可选备用）
# ---------------------------------------------------------------------------


def save_resume_docx(
    markdown_text: str,
    output_dir: str | Path,
    job_id: str,
    date_str: Optional[str] = None,
) -> Optional[Path]:
    """
    将 Markdown 简历导出为 DOCX 格式（python-docx）。

    DOCX 格式便于用户在 Word 中进一步编辑，作为 PDF 的补充产物。

    Returns:
        DOCX 文件路径，失败时返回 None
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx 未安装，无法生成 DOCX")
        return None

    try:
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        date_str = date_str or datetime.now().strftime("%Y%m%d")
        safe_id = re.sub(r"[^\w\-]", "_", job_id)
        docx_path = output_dir / f"{safe_id}_{date_str}.docx"

        doc = Document()
        _render_docx(doc, markdown_text)
        doc.save(str(docx_path))
        logger.info("简历 DOCX 已保存：%s", docx_path)
        return docx_path

    except Exception as e:
        logger.error("DOCX 生成失败：%s", e)
        return None


def _render_docx(doc, text: str) -> None:
    """按 Markdown 结构向 DOCX 文档写入内容。"""
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for line in text.splitlines():
        raw = line.rstrip()

        if not raw.strip():
            doc.add_paragraph()
            continue

        if re.match(r"^-{3,}$", raw.strip()):
            doc.add_paragraph("─" * 40)
            continue

        if raw.startswith("# "):
            p = doc.add_heading(raw[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if raw.startswith("## "):
            doc.add_heading(_strip_md(raw[3:].strip()), level=2)
            continue

        if raw.startswith("### "):
            doc.add_heading(_strip_md(raw[4:].strip()), level=3)
            continue

        if re.match(r"^[\-\*]\s+", raw):
            content = re.sub(r"^[\-\*]\s+", "", raw)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_md(p, _strip_md(content))
            continue

        if raw.startswith("> "):
            p = doc.add_paragraph(_strip_md(raw[2:]))
            p.style.font.italic = True
            continue

        p = doc.add_paragraph()
        _add_inline_md(p, raw)


def _add_inline_md(paragraph, text: str) -> None:
    """向 DOCX 段落写入内联 Markdown（处理加粗）。"""
    # 按 **bold** 分割
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(_strip_md(part))
